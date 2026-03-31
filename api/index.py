from fastapi import FastAPI, Response
from pydantic import BaseModel
import requests
import json
import uuid
import os
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document  # 新增：Word文档生成
from docx.shared import Pt  # 新增：字体大小设置
from docx.oxml.ns import qn  # 新增：中文字体设置
import io  # 新增：内存流

app = FastAPI()

# 解决跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 临时存储生成的内容（单机版，生产环境可改用数据库/Redis）
generated_content = {}

# 加载 Prompt 模板
try:
    with open("prompt_templates.json", "r", encoding="utf-8") as f:
        prompt_config = json.load(f)
    templates = prompt_config.get("templates", {})
except:
    templates = {}

# ====================== 替换为你的真实豆包 API Key ======================
API_KEY = "992f03a7-b58f-4850-8c86-c485b04e3ccd"
MODEL_ID = "ep-20260323110516-75srd"  # 替换为你的模型ID
# =======================================================================
API_URL = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"

class Item(BaseModel):
    type: str
    style: str
    len: str
    content: str

# 生成文书核心接口（新增存储生成内容的ID）
@app.post("/api/generate")
async def generate(item: Item):
    # 1. 校验API Key
    if not API_KEY or API_KEY == "你的豆包API密钥":
        return {"result": "❌ 错误：未配置有效豆包API Key，请在代码中替换API_KEY字段", "doc_id": ""}
    
    # 2. 校验文书类型
    if item.type not in templates:
        return {"result": f"❌ 错误：暂不支持「{item.type}」类型的文书生成", "doc_id": ""}
    
    # 3. 生成Prompt
    try:
        template = templates[item.type]["prompt"]
        style = "正式" if item.type == "述职报告" else item.style
        prompt = template.format(style=style, wordCount=item.len, content=item.content)
    except Exception as e:
        return {"result": f"❌ 模板解析错误：{str(e)}", "doc_id": ""}

    # 4. 调用豆包API
    try:
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": MODEL_ID,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 2000
        }

        resp = requests.post(API_URL, json=payload, headers=headers, timeout=30)
        result_data = resp.json()

        # 5. 解析结果并存储
        if resp.status_code == 200 and "choices" in result_data and result_data["choices"]:
            text = result_data["choices"][0]["message"]["content"].strip()
            doc_id = str(uuid.uuid4())  # 生成唯一文档ID
            generated_content[doc_id] = {
                "title": item.type,
                "content": text
            }
            return {"result": text, "doc_id": doc_id}  # 返回文档ID
        elif "error" in result_data:
            error_msg = result_data["error"]["message"]
            return {"result": f"❌ 豆包API错误：{error_msg}", "doc_id": ""}
        else:
            return {"result": "❌ 生成失败：AI未返回有效内容，请检查背景信息或重试", "doc_id": ""}
            
    except requests.exceptions.ConnectionError:
        return {"result": "❌ 网络错误：无法连接到豆包服务器，请稍后重试", "doc_id": ""}
    except requests.exceptions.Timeout:
        return {"result": "❌ 超时错误：请求豆包API超时（30秒），请重试", "doc_id": ""}
    except Exception as e:
        return {"result": f"❌ 系统异常：{str(e)}", "doc_id": ""}

# 新增：生成Word文档并下载的接口（极简稳定版）
@app.get("/api/download/{doc_id}")
async def download_doc(doc_id: str):
    # 校验文档ID是否存在
    if doc_id not in generated_content:
        return Response(content="文档不存在或已过期", status_code=404)
    
    # 获取生成的内容
    doc_data = generated_content[doc_id]
    title = doc_data["title"]
    content = doc_data["content"]

    try:
        # 1. 创建Word文档（移除复杂字体设置，避免报错）
        doc = Document()
        
        # 2. 添加标题（仅基础样式，不设置中文字体）
        doc.add_heading(title, level=1)
        
        # 3. 添加正文内容（基础样式）
        doc.add_paragraph(content)

        # 4. 保存到内存流
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)  # 重置流指针

        # 5. 返回下载响应
        return StreamingResponse(
            doc_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={title}.docx".encode("utf-8").decode("latin-1")
            }
        )
    except Exception as e:
        # 捕获所有异常，返回具体错误信息
        return Response(content=f"文档生成失败：{str(e)}", status_code=500)
        
# 获取场景参考样例
@app.get("/api/examples/{doc_type}")
async def get_example(doc_type: str):
    example = templates.get(doc_type, {}).get("example", "")
    return {"example": example}

# 首页返回前端页面
@app.get("/")
async def root():
    return FileResponse("index.html")